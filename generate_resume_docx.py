import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Helper to add clickable hyperlink to paragraph
def add_hyperlink(paragraph, text, url, color="1E3A8A", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_docx():
    docx_path = os.path.abspath("assets/Naveen_Mandula_Resume.docx")
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    
    doc = docx.Document()
    
    # 0.5-inch margins for data density
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Document Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark grey #1f2937

    PRIMARY_COLOR = RGBColor(0x1E, 0x3A, 0x8A) # Navy #1e3a8a
    MUTED_COLOR = RGBColor(0x4B, 0x55, 0x63) # Muted grey #4b5563

    # 1. Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("NAVEEN MANDULA")
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR

    # 2. Header Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(8)
    
    # Subtitle role
    run_sub = p_contact.add_run("Senior AI Architect | AIOps Engineer | AWS & Azure\n")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = MUTED_COLOR
    
    # Phone
    run_phone = p_contact.add_run("+1 (334) 454-0408  |  ")
    run_phone.font.name = 'Arial'
    run_phone.font.size = Pt(9)
    run_phone.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Email (Hyperlinked)
    add_hyperlink(p_contact, "nmandula0511@gmail.com", "mailto:nmandula0511@gmail.com", color="0284C7", underline=True)
    
    run_sep1 = p_contact.add_run("  |  ")
    run_sep1.font.name = 'Arial'
    run_sep1.font.size = Pt(9)
    
    # LinkedIn (Hyperlinked)
    add_hyperlink(p_contact, "linkedin", "https://www.linkedin.com/in/naveenmandula", color="0284C7", underline=True)
    
    run_sep2 = p_contact.add_run("  |  ")
    run_sep2.font.name = 'Arial'
    run_sep2.font.size = Pt(9)
    
    # GitHub (Hyperlinked)
    add_hyperlink(p_contact, "github", "https://github.com/nmandula0511", color="0284C7", underline=True)
    
    run_sep3 = p_contact.add_run("  |  ")
    run_sep3.font.name = 'Arial'
    run_sep3.font.size = Pt(9)
    
    # Portfolio (Hyperlinked)
    add_hyperlink(p_contact, "portfolio", "https://website-zeta-neon-17.vercel.app/", color="0284C7", underline=True)

    # Section Helper
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

    # 1. Professional Summary
    add_section_header("PROFESSIONAL SUMMARY")
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(6)
    p_summary.paragraph_format.line_spacing = 1.15
    run_sum = p_summary.add_run(
        "Highly accomplished Senior AI Architect and AIOps Engineer with over 10 years of experience "
        "designing, building, and optimizing production-grade intelligent platforms and cloud-native "
        "backend systems across telecommunications, logistics, finance, and industrial IoT domains. "
        "Proven expert in implementing scalable agentic AI networks, multi-agent orchestration frameworks, "
        "and robust event-driven streaming data pipelines. Demonstrated ability to bridge the gap between "
        "legacy Python/Linux environments and cutting-edge, high-availability AWS architectures while "
        "establishing rigorous cloud governance, performance optimization, and concrete cost reductions."
    )
    run_sum.font.name = 'Arial'
    run_sum.font.size = Pt(9.5)

    # 2. Technical Skills
    add_section_header("TECHNICAL SKILLS")
    skills = [
        ("AI Architecture & Agentic Systems: ", "Strands Framework (AWS-Native Agent Orchestration), Agent-to-Agent (A2A) Communication Protocols, Model Context Protocol (MCP) Gateways, Multi-Agent Systems, LangGraph, LangChain, Tool Calling, Planning/ReAct Agents, Memory Systems, AI Governance & Guardrails"),
        ("Generative AI & LLMs: ", "Amazon Bedrock Architecture, Amazon Nova PRO, GPT-4, Claude (Anthropic), Native S3 RAG Pipelines, Prompt Engineering, Hallucination Mitigation, Model Drift, RAGAS / DeepEval"),
        ("Cloud & Infrastructure: ", "AWS (Bedrock, SageMaker, Glue Jobs/ETL, Athena, Fargate, Lambda, EKS, S3 & S3 Tables, Step Functions, SNS/SQS, IAM), GCP (Vertex AI, BigQuery, Dataflow), Azure AI"),
        ("Databases & Data Streaming: ", "Apache Kafka (Real-Time Ingestion/Streaming), Amazon DynamoDB, Amazon Aurora (PostgreSQL Engine exclusively), Amazon Neptune (Graph DB/Digital Twin Modeling)"),
        ("DevOps & MLOps: ", "Terraform (Infrastructure as Code), CI/CD (GitHub Actions, GitLab Pipelines), Docker, Kubernetes (EKS/AKS), MLflow, OpenTelemetry, Prometheus, Grafana, CloudWatch"),
        ("Programming Languages: ", "Python (Expert), Python FastAPI, SQL, Scala, R, Go, Bash, YAML, JSON")
    ]
    for category, detail in skills:
        p_skill = doc.add_paragraph()
        p_skill.paragraph_format.space_after = Pt(3)
        run_cat = p_skill.add_run(category)
        run_cat.font.name = 'Arial'
        run_cat.font.size = Pt(9.5)
        run_cat.font.bold = True
        run_det = p_skill.add_run(detail)
        run_det.font.name = 'Arial'
        run_det.font.size = Pt(9.5)

    # 3. Professional Experience
    add_section_header("PROFESSIONAL EXPERIENCE")

    # Job 1: Charter
    p_job1_hdr = doc.add_paragraph()
    p_job1_hdr.paragraph_format.space_before = Pt(6)
    p_job1_hdr.paragraph_format.space_after = Pt(2)
    p_job1_hdr.paragraph_format.keep_with_next = True
    
    run_comp1 = p_job1_hdr.add_run("Charter Communications (Spectrum) ")
    run_comp1.font.name = 'Arial'
    run_comp1.font.size = Pt(9.5)
    run_comp1.font.bold = True
    
    run_loc1 = p_job1_hdr.add_run("|  Greenwood Village, CO")
    run_loc1.font.name = 'Arial'
    run_loc1.font.size = Pt(9.5)
    
    run_date1 = p_job1_hdr.add_run("\t\t\t\t\tFeb 2026 – Present")
    run_date1.font.name = 'Arial'
    run_date1.font.size = Pt(9.5)
    run_date1.font.bold = True

    p_job1_sub = doc.add_paragraph()
    p_job1_sub.paragraph_format.space_after = Pt(4)
    p_job1_sub.paragraph_format.keep_with_next = True
    run_title1 = p_job1_sub.add_run("Senior AWS Backend & AIOps Engineer")
    run_title1.font.name = 'Arial'
    run_title1.font.size = Pt(9)
    run_title1.font.italic = True
    run_title1.font.color.rgb = MUTED_COLOR

    charter_bullets = [
        "Core Platform Architecture: Architecting and delivering core AWS backend capabilities and scalable agentic AI models for the internal AIOps Platform, serving network operations groups to automate root cause analysis (RCA) and accelerate network incident remediation.",
        "AWS Data Pipeline Engineering: Engineered and optimized production AWS Glue ETL jobs to consume, clean, and pre-process high-volume telemetry data from Amazon S3, routing structured data directly into high-performance DynamoDB and Amazon Aurora data tiers.",
        "Agentic AI Refactoring: Re-architected a legacy 'vanilla Python' sub-agent codebase into a modular, production-grade Strands multi-agent layout running native Amazon Bedrock orchestration powered by the Amazon Nova PRO large language model.",
        "Agent-to-Agent (A2A) Protocols: Designed and deployed an asynchronous Agent-to-Agent (A2A) communication framework enabling the primary master orchestrator agent ('Paul') to dynamically discover and query domain-specific sub-agents via strict agent_card.json capability blueprints.",
        "Model Context Protocol (MCP) Integration: Replaced rigid sequential utility functions with a multi-threaded, parallelized Model Context Protocol (MCP) Gateway execution model, transforming business logic into reusable, discoverable tool-calling structures to reduce UI dashboard response latency.",
        "Production Code Standards: Upgraded system quality and statefulness across long-running, parallel user context boundaries by applying software engineering Factory Design Patterns and strict data validation contracts utilizing Pydantic models.",
        "API & Core Optimization: Streamlined backend performance by deprecating middle-tier Lambda wrappers and API Gateways, exposing REST invocation endpoints directly from the Agent Core layer to minimize request overhead and latency.",
        "Invincible Wi-Fi Real-Time Analytics: Maintained live, event-driven Apache Kafka data streaming pipelines capturing state triggers for 140k+ backup systems; implemented complex time-series correlation logic to flag anomalous loops where customer devices failed to roll back from 5G cellular backup to primary fiber paths after 90+ minutes.",
        "Cross-Domain Correlation Engine: Built data mechanics pairing cable modem status flags with backup unit telemetry, delivering instant visual root cause analysis (RCA) insights to engineers and preventing costly physical truck rolls.",
        "Infrastructure as Code: Controlled, provisioned, and scaled all platform infrastructure across DEVO testing spaces using Terraform scripts integrated with automated CI/CD pipelines."
    ]
    for bullet in charter_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2.5)
        parts = bullet.split(": ", 1)
        if len(parts) == 2:
            run_b1 = p_b.add_run(parts[0] + ": ")
            run_b1.font.bold = True
            run_b2 = p_b.add_run(parts[1])
        else:
            p_b.add_run(bullet)

    # Job 2: UPS
    p_job2_hdr = doc.add_paragraph()
    p_job2_hdr.paragraph_format.space_before = Pt(6)
    p_job2_hdr.paragraph_format.space_after = Pt(2)
    p_job2_hdr.paragraph_format.keep_with_next = True
    
    run_comp2 = p_job2_hdr.add_run("UPS (United Parcel Service) ")
    run_comp2.font.name = 'Arial'
    run_comp2.font.size = Pt(9.5)
    run_comp2.font.bold = True
    
    run_loc2 = p_job2_hdr.add_run("|  Atlanta, GA")
    run_loc2.font.name = 'Arial'
    run_loc2.font.size = Pt(9.5)
    
    run_date2 = p_job2_hdr.add_run("\t\t\t\t\t\tFeb 2022 – Feb 2026")
    run_date2.font.name = 'Arial'
    run_date2.font.size = Pt(9.5)
    run_date2.font.bold = True

    p_job2_sub = doc.add_paragraph()
    p_job2_sub.paragraph_format.space_after = Pt(4)
    p_job2_sub.paragraph_format.keep_with_next = True
    run_title2 = p_job2_sub.add_run("AI Engineer")
    run_title2.font.name = 'Arial'
    run_title2.font.size = Pt(9)
    run_title2.font.italic = True
    run_title2.font.color.rgb = MUTED_COLOR

    ups_bullets = [
        "Architected enterprise-scale Agentic AI platform with LangGraph-based multi-agent orchestration enabling intelligent automation across logistics, operations, and customer service workflows.",
        "Led end-to-end transition of AI solutions from PoC to production, standardizing reusable architectures for agentic workflows and RAG pipelines across multiple business units.",
        "Designed multi-agent orchestration frameworks with planner-executor patterns, dynamic tool routing, and stateful execution supporting complex enterprise decision-making.",
        "Built scalable multi-tenant AI platform on AWS Bedrock & SageMaker with secure onboarding, governance controls, and IAM-based access management.",
        "Developed advanced RAG pipelines integrating document ingestion, semantic chunking, embedding generation, and hybrid vector search (FAISS + Pinecone + BM25) to minimize hallucinations.",
        "Established AI governance frameworks including prompt validation, guardrails, audit logging, and compliance mechanisms ensuring responsible AI usage.",
        "Implemented LLM token optimization and cost governance strategies (context window optimization, caching) reducing inference costs while maintaining performance.",
        "Developed LLM observability solutions using OpenTelemetry and CloudWatch tracking latency, token usage, and throughput across distributed AI services.",
        "Built event-driven AI architectures using AWS SNS/SQS for asynchronous communication between agents, APIs, and backend services."
    ]
    for bullet in ups_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2.5)
        p_b.add_run(bullet)

    # Job 3: Visa
    p_job3_hdr = doc.add_paragraph()
    p_job3_hdr.paragraph_format.space_before = Pt(6)
    p_job3_hdr.paragraph_format.space_after = Pt(2)
    p_job3_hdr.paragraph_format.keep_with_next = True
    
    run_comp3 = p_job3_hdr.add_run("Visa Inc. ")
    run_comp3.font.name = 'Arial'
    run_comp3.font.size = Pt(9.5)
    run_comp3.font.bold = True
    
    run_loc3 = p_job3_hdr.add_run("|  Foster City, CA")
    run_loc3.font.name = 'Arial'
    run_loc3.font.size = Pt(9.5)
    
    run_date3 = p_job3_hdr.add_run("\t\t\t\t\t\t\tMay 2019 – Jan 2022")
    run_date3.font.name = 'Arial'
    run_date3.font.size = Pt(9.5)
    run_date3.font.bold = True

    p_job3_sub = doc.add_paragraph()
    p_job3_sub.paragraph_format.space_after = Pt(4)
    p_job3_sub.paragraph_format.keep_with_next = True
    run_title3 = p_job3_sub.add_run("Senior Data Scientist")
    run_title3.font.name = 'Arial'
    run_title3.font.size = Pt(9)
    run_title3.font.italic = True
    run_title3.font.color.rgb = MUTED_COLOR

    visa_bullets = [
        "Engineered real-time fraud detection platform using deep learning (neural networks) and ensemble methods (XGBoost, LightGBM/TensorFlow), achieving 88% detection accuracy and reducing false positives by 14% across billions of global payment transactions.",
        "Developed transaction authorization optimization models using ML and risk scoring algorithms, improving approval rates by 7% while maintaining robust fraud controls.",
        "Architected payment network analytics platform on GCP (BigQuery, Vertex AI, Dataflow) with distributed pipelines, reducing analytics processing time by 55%.",
        "Designed MLOps infrastructure using MLflow, Kubernetes, and Azure ML, monitoring 18+ production models across fraud detection, authorization, and risk management.",
        "Applied explainable AI techniques (SHAP/LIME) for transparency in risk and fraud models; developed PCI-compliant ML pipelines with data encryption and IAM governance."
    ]
    for bullet in visa_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2.5)
        p_b.add_run(bullet)

    # Job 4: GE
    p_job4_hdr = doc.add_paragraph()
    p_job4_hdr.paragraph_format.space_before = Pt(6)
    p_job4_hdr.paragraph_format.space_after = Pt(2)
    p_job4_hdr.paragraph_format.keep_with_next = True
    
    run_comp4 = p_job4_hdr.add_run("General Electric (GE) ")
    run_comp4.font.name = 'Arial'
    run_comp4.font.size = Pt(9.5)
    run_comp4.font.bold = True
    
    run_loc4 = p_job4_hdr.add_run("|  Boston, MA")
    run_loc4.font.name = 'Arial'
    run_loc4.font.size = Pt(9.5)
    
    run_date4 = p_job4_hdr.add_run("\t\t\t\t\t\t\tJul 2016 – Apr 2019")
    run_date4.font.name = 'Arial'
    run_date4.font.size = Pt(9.5)
    run_date4.font.bold = True

    p_job4_sub = doc.add_paragraph()
    p_job4_sub.paragraph_format.space_after = Pt(4)
    p_job4_sub.paragraph_format.keep_with_next = True
    run_title4 = p_job4_sub.add_run("Machine Learning Engineer")
    run_title4.font.name = 'Arial'
    run_title4.font.size = Pt(9)
    run_title4.font.italic = True
    run_title4.font.color.rgb = MUTED_COLOR

    ge_bullets = [
        "Built predictive maintenance platform for industrial equipment using ML (Random Forests, XGBoost) and IoT sensor analytics, predicting failures with 79% accuracy and reducing unplanned downtime by 11% across aviation, power, and renewable energy divisions.",
        "Developed energy production optimization models power plants using ML and physics-based simulations, improving plant efficiency by 6%.",
        "Architected supply chain demand forecasting models using ensemble methods and time-series analysis (LSTM, ARIMA), reducing stockouts by 14% across global markets.",
        "Architected AWS data infrastructure (EC2, Redshift, S3) for industrial IoT analytics processing sensor data from thousands of connected assets using distributed PySpark pipelines.",
        "Designed LLM-assisted diagnostic workflows for engineering report generation, anomaly summarization, and insights delivery.",
        "Automated infrastructure provisioning using Terraform and AWS IaC patterns under secure VPC isolation."
    ]
    for bullet in ge_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2.5)
        p_b.add_run(bullet)

    # Job 5: Abbott
    p_job5_hdr = doc.add_paragraph()
    p_job5_hdr.paragraph_format.space_before = Pt(6)
    p_job5_hdr.paragraph_format.space_after = Pt(2)
    p_job5_hdr.paragraph_format.keep_with_next = True
    
    run_comp5 = p_job5_hdr.add_run("Abbott Laboratories ")
    run_comp5.font.name = 'Arial'
    run_comp5.font.size = Pt(9.5)
    run_comp5.font.bold = True
    
    run_loc5 = p_job5_hdr.add_run("|  Abbott Park, IL")
    run_loc5.font.name = 'Arial'
    run_loc5.font.size = Pt(9.5)
    
    run_date5 = p_job5_hdr.add_run("\t\t\t\t\t\t\tAug 2013 – Jun 2016")
    run_date5.font.name = 'Arial'
    run_date5.font.size = Pt(9.5)
    run_date5.font.bold = True

    p_job5_sub = doc.add_paragraph()
    p_job5_sub.paragraph_format.space_after = Pt(4)
    p_job5_sub.paragraph_format.keep_with_next = True
    run_title5 = p_job5_sub.add_run("Data Scientist")
    run_title5.font.name = 'Arial'
    run_title5.font.size = Pt(9)
    run_title5.font.italic = True
    run_title5.font.color.rgb = MUTED_COLOR

    abbott_bullets = [
        "Developed medical device quality prediction models using ML (Random Forests, XGBoost) on manufacturing data, reducing defect rates by 9% and improving product quality across diagnostics and medical devices.",
        "Built demand forecasting models for pharmaceuticals and medical devices using ARIMA and regression techniques, achieving 86% accuracy to optimize production planning and inventory management.",
        "Engineered clinical trial analytics using statistical modeling and survival analysis on patient data from cardiovascular and diabetes studies, accelerating regulatory submissions.",
        "Architected supply chain optimization models using linear programming and simulation, reducing disruptions by 8% across pharmaceutical and device manufacturing."
    ]
    for bullet in abbott_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2.5)
        p_b.add_run(bullet)

    # 4. Education
    add_section_header("EDUCATION")
    p_edu1 = doc.add_paragraph()
    p_edu1.paragraph_format.space_after = Pt(2)
    run_e1 = p_edu1.add_run("Master of Science in Computer Science ")
    run_e1.font.name = 'Arial'
    run_e1.font.bold = True
    run_e2 = p_edu1.add_run("— Auburn University (2012 – 2013)")
    run_e2.font.name = 'Arial'

    p_edu2 = doc.add_paragraph()
    p_edu2.paragraph_format.space_after = Pt(2)
    run_e3 = p_edu2.add_run("Bachelor of Technology (B.Tech), Computer Science & Engineering ")
    run_e3.font.name = 'Arial'
    run_e3.font.bold = True
    run_e4 = p_edu2.add_run("— Jawaharlal Nehru Technological University (JNTU)  |  Hyderabad, India")
    run_e4.font.name = 'Arial'

    doc.save(docx_path)
    print("DOCX generated successfully.")

if __name__ == "__main__":
    generate_docx()
